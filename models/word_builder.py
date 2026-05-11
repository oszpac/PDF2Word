from docx import Document
from docx.shared import Pt, Inches
from utils.font_utils import set_cn_font
from config import FONT_SIZE_DEFAULT


class WordBuilder:

    def __init__(self, pages_data, pdf_model, progress_callback=None, log_callback=None):
        self.pages_data = pages_data
        self.pdf_model = pdf_model
        self.progress_callback = progress_callback
        self.log_callback = log_callback
        self._doc = Document()

    def build(self):
        self._log('开始构建Word文档...')
        self._setup_document()

        total = len(self.pages_data)
        for i, page_data in enumerate(self.pages_data):
            if self.progress_callback:
                self.progress_callback(i + 1, total, f'正在写入Word第 {i + 1}/{total} 页')
            is_last = (i == total - 1)
            self._add_page(page_data, add_page_break=not is_last)

        self._log(f'Word文档构建完成，共 {total} 页')
        return self._doc

    def save(self, output_path):
        try:
            self._doc.save(output_path)
            self._log(f'已保存到: {output_path}')
        except PermissionError:
            import tempfile
            import os
            fallback_dir = os.path.join(tempfile.gettempdir(), 'PDF2Word')
            os.makedirs(fallback_dir, exist_ok=True)
            fallback_path = os.path.join(fallback_dir, os.path.basename(output_path))
            counter = 1
            while os.path.exists(fallback_path):
                name, ext = os.path.splitext(os.path.basename(output_path))
                fallback_path = os.path.join(fallback_dir, f'{name}({counter}){ext}')
                counter += 1
            self._doc.save(fallback_path)
            self._log(f'原路径无写入权限，已保存到: {fallback_path}')
            return fallback_path
        return output_path

    def _setup_document(self):
        style = self._doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(FONT_SIZE_DEFAULT)

        section = self._doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)

    def _add_page(self, page_data, add_page_break=True):
        blocks = page_data.get('blocks', [])
        if not blocks:
            return

        last_bottom = -1
        GAP_SMALL = 3
        GAP_NORMAL = 8
        GAP_LARGE = 18

        for block in blocks:
            block_type = block.get('type', 'text')

            if block_type == 'text':
                if last_bottom >= 0:
                    top_gap = block['y'] - last_bottom
                    if top_gap > GAP_SMALL:
                        spacer = self._doc.add_paragraph()
                        spacer.paragraph_format.space_before = Pt(0)
                        spacer.paragraph_format.space_after = Pt(0)
                        if top_gap > 100:
                            gap_pt = GAP_LARGE
                        elif top_gap > 30:
                            gap_pt = GAP_NORMAL
                        else:
                            gap_pt = GAP_SMALL
                        spacer.paragraph_format.line_spacing = Pt(gap_pt)
                        run = spacer.add_run('')
                        run.font.size = Pt(1)

                para = self._doc.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = Pt(max(block.get('height', 12), 12))

                text = block.get('text', '')
                run = para.add_run(text)

                spans = block.get('spans', [])
                if spans:
                    first_span = spans[0]
                    font_size = first_span.get('size', FONT_SIZE_DEFAULT)
                else:
                    font_size = FONT_SIZE_DEFAULT

                set_cn_font(run, font_size=Pt(font_size))

                last_bottom = block['y'] + block.get('height', 12)

            elif block_type == 'image':
                self._add_image_block(block)

        if add_page_break:
            self._doc.add_page_break()

    def _add_image_block(self, block):
        try:
            page_index = block.get('page_index')
            if page_index is not None:
                page = self.pdf_model.doc[page_index]
                images = page.get_images(full=True)
                img_index = block.get('image_index', -1)
                if 0 <= img_index < len(images):
                    xref = images[img_index][0]
                    base_image = self.pdf_model.doc.extract_image(xref)
                    image_bytes = base_image['image']
                    import io
                    image_stream = io.BytesIO(image_bytes)
                    width_inches = min(block.get('width', 200) / 72, 6.0)
                    self._doc.add_picture(image_stream, width=Inches(width_inches))
        except Exception as e:
            self._log(f'嵌入图片失败: {e}')

    def get_document(self):
        return self._doc

    def _log(self, message):
        if self.log_callback:
            self.log_callback(message)
